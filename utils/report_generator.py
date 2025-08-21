import io
import tempfile
import os
from datetime import datetime
from typing import Dict, List, Any
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_combined_chart(test_results: Dict[str, Any]) -> str:
    """
    Cria gráfico combinado de comparação entre tipos de teste usando subplots (um por métrica).
    
    Args:
        test_results: Dicionário com resultados das avaliações por tipo de teste
    
    Returns:
        str: Caminho para o arquivo de imagem temporário
    """
    # Determinar todas as métricas únicas disponíveis
    all_evaluations = set()
    for test_type in test_results:
        result_data = test_results[test_type]
        all_evaluations.update(result_data['evaluations'])
    
    all_evaluations = sorted(list(all_evaluations))
    
    if not all_evaluations:
        # Fallback para caso sem avaliações
        fig = go.Figure()
        fig.add_annotation(text="Nenhuma avaliação encontrada", xref="paper", yref="paper", x=0.5, y=0.5)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        fig.write_image(temp_file.name, format='png', engine='kaleido')
        temp_file.close()
        return temp_file.name
    
    # Preparar nomes limpos das avaliações
    clean_evaluation_names = []
    for eval_config in all_evaluations:
        clean_name = eval_config.replace("Factual Correctness - ", "FC - ").replace("Semantic Similarity - ", "SS - ")
        clean_evaluation_names.append(clean_name)
    
    # Calcular layout dos subplots - um subplot por métrica
    n_metrics = len(all_evaluations)
    cols = min(2, n_metrics)  # Máximo 2 colunas
    rows = (n_metrics + cols - 1) // cols  # Calcular número de linhas necessárias
    
    fig = make_subplots(
        rows=rows, 
        cols=cols,
        subplot_titles=clean_evaluation_names,
        vertical_spacing=0.15,
        horizontal_spacing=0.15
    )
    
    colors_list = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22']
    
    # Para cada métrica, criar um subplot
    for idx, eval_config in enumerate(all_evaluations):
        row = (idx // cols) + 1
        col = (idx % cols) + 1
        
        test_type_names = []
        scores = []
        
        # Para cada tipo de teste, calcular a média da métrica
        for test_type in test_results.keys():
            result_data = test_results[test_type]
            
            if eval_config in result_data['evaluations']:
                metric_scores = []
                for result in result_data['results']:
                    if isinstance(result.get(eval_config), (int, float)):
                        metric_scores.append(result[eval_config])
                
                if metric_scores:
                    test_type_names.append(test_type)
                    scores.append(sum(metric_scores) / len(metric_scores))
        
        if test_type_names and scores:
            fig.add_trace(
                go.Bar(
                    x=test_type_names,
                    y=scores,
                    name=clean_evaluation_names[idx],
                    marker_color=[colors_list[i % len(colors_list)] for i in range(len(test_type_names))],
                    text=[f"{score:.3f}" for score in scores],
                    textposition='outside',
                    showlegend=False
                ),
                row=row, col=col
            )
            
            # Configurar eixos para cada subplot
            fig.update_yaxes(range=[0, 1.1], row=row, col=col)
            fig.update_xaxes(tickangle=45, row=row, col=col)
    
    fig.update_layout(
        title="Comparação Detalhada: Avaliações por Tipo de Teste",
        height=400 * rows,  # Altura dinâmica baseada no número de linhas
        width=800,
        showlegend=False
    )
    
    # Salvar como imagem temporária
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    fig.write_image(temp_file.name, format='png', engine='kaleido')
    temp_file.close()
    
    return temp_file.name


def create_individual_chart(test_type: str, result_data: Dict[str, Any]) -> str:
    """
    Cria gráfico individual para um tipo de teste específico.
    
    Args:
        test_type: Nome do tipo de teste
        result_data: Dados dos resultados
    
    Returns:
        str: Caminho para o arquivo de imagem temporário
    """
    all_results = result_data['results']
    selected_evaluations = result_data['evaluations']
    
    # Preparar dados para gráfico de barras com médias
    bar_data = []
    
    for eval_config in selected_evaluations:
        scores = []
        for result in all_results:
            if isinstance(result[eval_config], (int, float)):
                scores.append(result[eval_config])
        
        if scores:  # Se há scores válidos
            # Clean up the evaluation name for display
            clean_name = eval_config.replace("Factual Correctness - ", "FC - ").replace("Semantic Similarity - ", "SS - ")
            media_scores = sum(scores) / len(scores)
            bar_data.append({
                'evaluation': clean_name,
                'media': media_scores
            })
    
    # Criar gráfico de barras
    fig = go.Figure()
    
    colors_list = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728', '#9467bd', '#8c564b', '#e377c2', '#7f7f7f', '#bcbd22']
    
    evaluations = [data['evaluation'] for data in bar_data]
    medias = [data['media'] for data in bar_data]
    
    fig.add_trace(go.Bar(
        x=evaluations,
        y=medias,
        marker_color=[colors_list[i % len(colors_list)] for i in range(len(evaluations))],
        text=[f"{media:.3f}" for media in medias],
        textposition='outside'
    ))
    
    fig.update_layout(
        title=f"Média dos Scores - {test_type}",
        xaxis_title="Tipo de Avaliação",
        yaxis_title="Score Médio",
        yaxis=dict(range=[0, max(1, max(medias) * 1.1) if medias else 1]),
        height=500,
        width=800,
        showlegend=False
    )
    
    # Salvar como imagem temporária
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
    fig.write_image(temp_file.name, format='png', engine='kaleido')
    temp_file.close()
    
    return temp_file.name


def generate_pdf_report(test_results: Dict[str, Any], test_types: Dict[str, Any]) -> bytes:
    """
    Gera um relatório em PDF dos resultados das avaliações.
    
    Args:
        test_results: Dicionário com resultados das avaliações por tipo de teste
        test_types: Dicionário com informações dos tipos de teste
    
    Returns:
        bytes: Conteúdo do PDF gerado
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    temp_files = []  # Lista para rastrear arquivos temporários
    
    try:
        # Título do relatório
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.darkblue
        )
        
        story.append(Paragraph("Relatório de Comparação de Textos RAGAS", title_style))
        story.append(Paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Resumo geral
        story.append(Paragraph("Resumo Geral", styles['Heading2']))
        story.append(Paragraph(f"Tipos de teste avaliados: {len(test_results)}", styles['Normal']))
        
        total_iterations = sum(len(data['results']) for data in test_results.values())
        story.append(Paragraph(f"Total de iterações: {total_iterations}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Gráfico de comparação geral (se há mais de um tipo de teste)
        if len(test_results) > 1:
            story.append(Paragraph("Comparação Geral entre Tipos de Teste", styles['Heading2']))
            
            try:
                combined_chart_path = create_combined_chart(test_results)
                temp_files.append(combined_chart_path)
                
                # Determinar altura baseada no número de métricas
                all_evaluations = set()
                for result_data in test_results.values():
                    all_evaluations.update(result_data['evaluations'])
                n_metrics = len(all_evaluations)
                cols = min(2, n_metrics)
                rows = (n_metrics + cols - 1) // cols
                chart_height = min(3 * rows, 8)  # Altura máxima de 8 inches
                
                # Adicionar imagem ao PDF
                img = Image(combined_chart_path, width=6*inch, height=chart_height*inch)
                story.append(img)
                story.append(Spacer(1, 20))
                
            except Exception as e:
                story.append(Paragraph(f"Erro ao gerar gráfico de comparação: {str(e)}", styles['Normal']))
                story.append(Spacer(1, 20))
        
        # Para cada tipo de teste
        for test_type, result_data in test_results.items():
            story.append(Paragraph(f"Tipo de Teste: {test_type}", styles['Heading2']))
            
            # Informações do tipo de teste
            if test_type in test_types:
                created_at = test_types[test_type]['created_at']
                story.append(Paragraph(f"Criado em: {created_at}", styles['Normal']))
            
            story.append(Paragraph(f"Timestamp da avaliação: {result_data['timestamp']}", styles['Normal']))
            story.append(Paragraph(f"Iterações avaliadas: {len(result_data['results'])}", styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Estatísticas por avaliação
            story.append(Paragraph("Resumo Estatístico", styles['Heading3']))
            
            # Preparar dados da tabela de estatísticas
            stats_data = [['Avaliação', 'Média', 'Mediana', 'Mínimo', 'Máximo', 'Desvio Padrão', 'Iterações']]
            
            for eval_config in result_data['evaluations']:
                scores = []
                for result in result_data['results']:
                    if isinstance(result.get(eval_config), (int, float)):
                        scores.append(result[eval_config])
                
                if scores:
                    clean_name = eval_config.replace("Factual Correctness - ", "FC - ").replace("Semantic Similarity - ", "SS - ")
                    mean_score = sum(scores) / len(scores)
                    median_score = sorted(scores)[len(scores)//2]
                    min_score = min(scores)
                    max_score = max(scores)
                    std_dev = (sum([(x - mean_score)**2 for x in scores]) / len(scores))**0.5
                    
                    stats_data.append([
                        clean_name,
                        f"{mean_score:.3f}",
                        f"{median_score:.3f}",
                        f"{min_score:.3f}",
                        f"{max_score:.3f}",
                        f"{std_dev:.3f}",
                        str(len(scores))
                    ])
            
            # Criar tabela de estatísticas
            stats_table = Table(stats_data)
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(stats_table)
            story.append(Spacer(1, 20))
            
            # Gráfico individual para este tipo de teste
            story.append(Paragraph("Visualização Gráfica", styles['Heading3']))
            
            try:
                individual_chart_path = create_individual_chart(test_type, result_data)
                temp_files.append(individual_chart_path)
                
                # Adicionar imagem ao PDF
                img = Image(individual_chart_path, width=6*inch, height=3.75*inch)
                story.append(img)
                story.append(Spacer(1, 20))
                
            except Exception as e:
                story.append(Paragraph(f"Erro ao gerar gráfico: {str(e)}", styles['Normal']))
                story.append(Spacer(1, 20))
            
            # Tabela de resultados detalhados
            story.append(Paragraph("Resultados Detalhados por Iteração", styles['Heading3']))
            
            # Preparar dados da tabela de resultados
            results_data = [['Iteração'] + [eval_config.replace("Factual Correctness - ", "FC - ").replace("Semantic Similarity - ", "SS - ") 
                                           for eval_config in result_data['evaluations']]]
            
            for result in result_data['results']:
                row = [result['iteration']]
                for eval_config in result_data['evaluations']:
                    value = result.get(eval_config)
                    if isinstance(value, (int, float)):
                        row.append(f"{value:.3f}")
                    else:
                        row.append(str(value))
                results_data.append(row)
            
            # Criar tabela de resultados
            results_table = Table(results_data)
            results_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgrey),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(results_table)
            
            # Quebra de página entre tipos de teste (exceto o último)
            if test_type != list(test_results.keys())[-1]:
                story.append(PageBreak())
            else:
                story.append(Spacer(1, 20))
        
        # Rodapé
        story.append(Spacer(1, 20))
        story.append(Paragraph("Powered by RAGAS - Factual Correctness Evaluation", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
        
    finally:
        # Limpar arquivos temporários
        for temp_file in temp_files:
            try:
                os.unlink(temp_file)
            except Exception:
                pass


def generate_docx_report(test_results: Dict[str, Any], test_types: Dict[str, Any]) -> bytes:
    """
    Gera um relatório em DOCX dos resultados das avaliações.
    
    Args:
        test_results: Dicionário com resultados das avaliações por tipo de teste
        test_types: Dicionário com informações dos tipos de teste
    
    Returns:
        bytes: Conteúdo do DOCX gerado
    """
    doc = Document()
    temp_files = []  # Lista para rastrear arquivos temporários
    
    try:
        # Título do relatório
        title = doc.add_heading('Relatório de Comparação de Textos RAGAS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data de geração
        date_para = doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Resumo geral
        doc.add_heading('Resumo Geral', level=1)
        doc.add_paragraph(f'Tipos de teste avaliados: {len(test_results)}')
        
        total_iterations = sum(len(data['results']) for data in test_results.values())
        doc.add_paragraph(f'Total de iterações: {total_iterations}')
        
        # Gráfico de comparação geral (se há mais de um tipo de teste)
        if len(test_results) > 1:
            doc.add_heading('Comparação Geral entre Tipos de Teste', level=1)
            
            try:
                combined_chart_path = create_combined_chart(test_results)
                temp_files.append(combined_chart_path)
                
                # Determinar altura baseada no número de métricas
                all_evaluations = set()
                for result_data in test_results.values():
                    all_evaluations.update(result_data['evaluations'])
                n_metrics = len(all_evaluations)
                cols = min(2, n_metrics)
                rows = (n_metrics + cols - 1) // cols
                chart_height = min(3 * rows, 8)  # Altura máxima de 8 inches
                
                # Adicionar imagem ao DOCX
                doc.add_picture(combined_chart_path, width=Inches(6), height=Inches(chart_height))
                
            except Exception as e:
                doc.add_paragraph(f'Erro ao gerar gráfico de comparação: {str(e)}')
        
        # Para cada tipo de teste
        for test_type, result_data in test_results.items():
            doc.add_heading(f'Tipo de Teste: {test_type}', level=1)
            
            # Informações do tipo de teste
            if test_type in test_types:
                created_at = test_types[test_type]['created_at']
                doc.add_paragraph(f'Criado em: {created_at}')
            
            doc.add_paragraph(f'Timestamp da avaliação: {result_data["timestamp"]}')
            doc.add_paragraph(f'Iterações avaliadas: {len(result_data["results"])}')
            
            # Estatísticas por avaliação
            doc.add_heading('Resumo Estatístico', level=2)
            
            # Criar tabela de estatísticas
            stats_table = doc.add_table(rows=1, cols=7)
            stats_table.style = 'Table Grid'
            hdr_cells = stats_table.rows[0].cells
            headers = ['Avaliação', 'Média', 'Mediana', 'Mínimo', 'Máximo', 'Desvio Padrão', 'Iterações']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
            
            for eval_config in result_data['evaluations']:
                scores = []
                for result in result_data['results']:
                    if isinstance(result.get(eval_config), (int, float)):
                        scores.append(result[eval_config])
                
                if scores:
                    clean_name = eval_config.replace("Factual Correctness - ", "FC - ").replace("Semantic Similarity - ", "SS - ")
                    mean_score = sum(scores) / len(scores)
                    median_score = sorted(scores)[len(scores)//2]
                    min_score = min(scores)
                    max_score = max(scores)
                    std_dev = (sum([(x - mean_score)**2 for x in scores]) / len(scores))**0.5
                    
                    row_cells = stats_table.add_row().cells
                    row_cells[0].text = clean_name
                    row_cells[1].text = f"{mean_score:.3f}"
                    row_cells[2].text = f"{median_score:.3f}"
                    row_cells[3].text = f"{min_score:.3f}"
                    row_cells[4].text = f"{max_score:.3f}"
                    row_cells[5].text = f"{std_dev:.3f}"
                    row_cells[6].text = str(len(scores))
            
            # Gráfico individual para este tipo de teste
            doc.add_heading('Visualização Gráfica', level=2)
            
            try:
                individual_chart_path = create_individual_chart(test_type, result_data)
                temp_files.append(individual_chart_path)
                
                # Adicionar imagem ao DOCX
                doc.add_picture(individual_chart_path, width=Inches(6))
                
            except Exception as e:
                doc.add_paragraph(f'Erro ao gerar gráfico: {str(e)}')
            
            # Resultados detalhados
            doc.add_heading('Resultados Detalhados por Iteração', level=2)
            
            # Criar tabela de resultados
            num_cols = 1 + len(result_data['evaluations'])
            results_table = doc.add_table(rows=1, cols=num_cols)
            results_table.style = 'Table Grid'
            
            # Cabeçalhos
            hdr_cells = results_table.rows[0].cells
            hdr_cells[0].text = 'Iteração'
            for i, eval_config in enumerate(result_data['evaluations']):
                clean_name = eval_config.replace("Factual Correctness - ", "FC - ").replace("Semantic Similarity - ", "SS - ")
                hdr_cells[i + 1].text = clean_name
            
            # Dados
            for result in result_data['results']:
                row_cells = results_table.add_row().cells
                row_cells[0].text = result['iteration']
                
                for i, eval_config in enumerate(result_data['evaluations']):
                    value = result.get(eval_config)
                    if isinstance(value, (int, float)):
                        row_cells[i + 1].text = f"{value:.3f}"
                    else:
                        row_cells[i + 1].text = str(value)
        
            # Quebra de página entre tipos de teste (exceto o último)
            if test_type != list(test_results.keys())[-1]:
                doc.add_page_break()
        
        # Rodapé
        doc.add_paragraph()
        footer_para = doc.add_paragraph('Powered by RAGAS - Factual Correctness Evaluation')
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Salvar em buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    finally:
        # Limpar arquivos temporários
        for temp_file in temp_files:
            try:
                os.unlink(temp_file)
            except Exception:
                pass